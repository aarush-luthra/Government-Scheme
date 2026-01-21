
import sys
import os
import re
from docx import Document
from docx.shared import Inches

# Add project root to path
sys.path.append(os.getcwd())

TESTS_DIR = os.path.join(os.getcwd(), "tests")

def convert_md_to_docx(filename_base):
    md_file = os.path.join(TESTS_DIR, f"{filename_base}.md")
    docx_file = os.path.join(TESTS_DIR, f"{filename_base}.docx")
    
    if not os.path.exists(md_file):
        print(f"Error: Markdown file not found at {md_file}")
        return

    print(f"Converting {md_file} -> {docx_file} ...")
    doc = Document()
    
    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # Simple Markdown Parser
    table_lines = []
    
    # Pre-process lines to handle tables (grouping)
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue

        # Check for table block (starts with |)
        if line.startswith("|"):
            table_lines.append(line)
            # Look ahead for more table lines
            while i + 1 < len(lines) and lines[i+1].strip().startswith("|"):
                i += 1
                table_lines.append(lines[i].strip())
            
            # Process the collected table
            if len(table_lines) >= 3: # Header + Separator + At least 1 row
                # Extract headers
                headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
                # Skip separator line (table_lines[1])
                
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Write Header
                hdr_cells = table.rows[0].cells
                for idx, h in enumerate(headers):
                    hdr_cells[idx].text = h
                
                # Write Data Rows
                for row_line in table_lines[2:]:
                    row_data = [d.strip() for d in row_line.split('|') if d.strip()]
                    # Handle row length mismatch just in case
                    if len(row_data) != len(headers):
                         # Try to fit or skip? Let's just fit what we can
                         pass
                    
                    row_cells = table.add_row().cells
                    for idx, d in enumerate(row_data):
                        if idx < len(row_cells):
                            row_cells[idx].text = d
            
            table_lines = [] # Reset
        
        # 1. Headers
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        
        # 2. Images ![Alt](Path)
        elif line.startswith("!["):
            # Extract path using regex
            match = re.search(r'\!\[.*?\]\((.*?)\)', line)
            if match:
                img_rel_path = match.group(1)
                # Cleaning path just in case
                if img_rel_path.startswith("./"):
                    img_rel_path = img_rel_path[2:]
                
                # Construct full path
                # Expecting path like "plots/filename.png" or "filename.png"
                img_path = os.path.join(TESTS_DIR, img_rel_path)
                
                if os.path.exists(img_path):
                    print(f"  Adding image: {img_path}")
                    try:
                        doc.add_picture(img_path, width=Inches(6))
                    except Exception as e:
                        print(f"  Failed to add image {img_path}: {e}")
                        doc.add_paragraph(f"[Image: {img_rel_path}]")
                else:
                    print(f"  Image not found: {img_path}")
                    doc.add_paragraph(f"[Missing Image: {img_rel_path}]")
        
        # 3. Code Blocks (simple detection)
        elif line.startswith("```"):
             # Just skip the fence markers, maybe set a flag for code style later
             # For now, just ignoring the fence line to avoid printing "```"
             pass

        # 4. Blockquotes
        elif line.startswith("> "):
            doc.add_paragraph(line[2:], style='Quote')

        # 5. Normal Text
        else:
            doc.add_paragraph(line)
        
        i += 1

    doc.save(docx_file)
    print(f"Successfully created: {docx_file}\n")

if __name__ == "__main__":
    convert_md_to_docx("journal_paper")
    convert_md_to_docx("benchmark_data_report")
